from rest_framework import status
from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import MultiPartParser
from rest_framework.response import Response
from docx import Document
from PyPDF2 import PdfReader

from .models import ParsedDocument
from .serializers import ParsedDocumentSerializer, ParseUploadSerializer


@api_view(['POST'])
@parser_classes([MultiPartParser])
def parse_document(request):
    serializer = ParseUploadSerializer(data=request.data)
    if not serializer.is_valid():
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    uploaded_file = serializer.validated_data['file']
    filename = uploaded_file.name
    file_ext = filename.lower().split('.')[-1]
    file_size = uploaded_file.size

    extracted_text = ''
    page_count = None

    try:
        if file_ext == 'pdf':
            reader = PdfReader(uploaded_file)
            page_count = len(reader.pages)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    extracted_text += text + '\n'

        elif file_ext == 'docx':
            doc = Document(uploaded_file)
            paragraphs = []
            for paragraph in doc.paragraphs:
                paragraphs.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.append(cell.text)
            extracted_text = '\n'.join(paragraphs)

    except Exception as e:
        return Response(
            {'error': f'Failed to parse document: {str(e)}'},
            status=status.HTTP_400_BAD_REQUEST
        )

    uploaded_file.seek(0)

    parsed_doc = ParsedDocument.objects.create(
        original_filename=filename,
        file_type=file_ext.upper(),
        file_size=file_size,
        page_count=page_count,
        extracted_text=extracted_text.strip(),
        original_file=uploaded_file
    )

    response_serializer = ParsedDocumentSerializer(parsed_doc)
    return Response(response_serializer.data, status=status.HTTP_201_CREATED)


@api_view(['GET'])
def get_parsed_document(request, pk):
    try:
        parsed_doc = ParsedDocument.objects.get(pk=pk)
    except ParsedDocument.DoesNotExist:
        return Response(
            {'error': 'Parsed document not found.'},
            status=status.HTTP_404_NOT_FOUND
        )

    serializer = ParsedDocumentSerializer(parsed_doc)
    return Response(serializer.data)
