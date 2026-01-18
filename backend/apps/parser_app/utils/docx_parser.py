from __future__ import annotations
import io
import math
from typing import List, Dict, Any, Optional
from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.oxml.ns import qn
from PIL import Image as PILImage
from .elements import (
    make_text_element, make_table_element, make_image_element,
    make_signature_element, is_likely_signature, make_divider_element, generate_id
)

# === ТОЧНЫЕ КОНСТАНТЫ WORD ===
PT_TO_PX = 1.3333333333333333  # 96 DPI / 72 DPI (pt to px)
MM_TO_PX = 3.937007874015748   # 96 DPI / 25.4 mm

# A4 при 96 DPI
PAGE_WIDTH = int(210 * MM_TO_PX)      # 794px
PAGE_HEIGHT = int(297 * MM_TO_PX)     # 1123px

# Стандартные поля Word (2.54 см)
DEFAULT_MARGIN = int(25.4 * MM_TO_PX)  # 96px
LEFT_MARGIN = DEFAULT_MARGIN
RIGHT_MARGIN = DEFAULT_MARGIN
TOP_MARGIN = DEFAULT_MARGIN
BOTTOM_MARGIN = DEFAULT_MARGIN

CONTENT_WIDTH = PAGE_WIDTH - (LEFT_MARGIN + RIGHT_MARGIN)

def parse_docx(file_obj) -> Dict[str, Any]:
    """Главная функция парсинга с АБСОЛЮТНЫМ позиционированием"""
    doc = Document(file_obj)
    elements: List[Dict[str, Any]] = []
    plain_chunks: List[str] = []
    y_offset = TOP_MARGIN
    
    # Загрузка реальных полей документа
    _load_document_margins(doc)
    
    # Кэш стилей
    style_cache = _build_style_cache(doc)
    
    # Карта изображений
    images_map = {}
    try:
        for rel in doc.part.rels.values():
            if rel.target_ref and ("image" in rel.target_ref or "media" in rel.target_ref):
                images_map[rel.rId] = rel.target_part.blob
    except Exception:
        pass

    # === ОСНОВНОЙ ЦИКЛ ПАРСИНГА ===
    for kind, block, section_idx in _iter_all_blocks(doc):
        if kind == 'paragraph':
            # Разрыв страницы
            if _has_page_break(block):
                y_offset = TOP_MARGIN
                plain_chunks.append("[Page Break]")
            
            # Форматирование параграфа
            fmt = _get_paragraph_format(block, style_cache)
            
            # Отступ ПЕРЕД
            y_offset += fmt.space_before
            
            # Разделитель (граница)
            if fmt.has_border_bottom:
                elements.append(make_divider_element(
                    x=LEFT_MARGIN, y=y_offset, width=CONTENT_WIDTH, thickness=max(1, fmt.border_width)
                ))
                y_offset += max(2, fmt.border_width) + 6
            
            # === ИЗОБРАЖЕНИЯ В ПАРАГРАФЕ ===
            blips = block._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
            has_image = False
            for blip in blips:
                embed = blip.get(qn('r:embed'))
                if embed and embed in images_map:
                    has_image = True
                    # Изображения выравниваются по левому краю с отступом параграфа
                    img_x = LEFT_MARGIN + fmt.left_indent
                    el, h_added = _process_image(images_map[embed], img_x, y_offset)
                    if el:
                        elements.append(el)
                        y_offset += h_added + 10  # Отступ под изображением
                        plain_chunks.append("[Image]")
            
            # === ТЕКСТ ===
            text_content = block.text.strip()
            if text_content:
                # X с учетом всех отступов
                final_x = LEFT_MARGIN + fmt.left_indent + fmt.first_line_indent
                eff_width = max(50, CONTENT_WIDTH - fmt.left_indent - fmt.right_indent)
                
                # ТОЧНАЯ ВЫСОТА
                total_h = _calculate_text_height_exact(text_content, fmt.size, eff_width, fmt.line_spacing)
                
                elements.append(make_text_element(
                    x=final_x, y=y_offset, width=eff_width, height=total_h,
                    content=text_content, font=fmt.font, size=fmt.size,
                    bold=fmt.bold, italic=fmt.italic, color=fmt.color, align=fmt.align
                ))
                y_offset += total_h
                plain_chunks.append(text_content)
            
            # Отступ ПОСЛЕ
            y_offset += fmt.space_after
            
            # Минимальный интервал для пустых параграфов
            if not text_content and not has_image and not fmt.has_border_bottom:
                y_offset += 8
            
            # Контроль высоты страницы
            if y_offset > PAGE_HEIGHT - BOTTOM_MARGIN - 40:
                y_offset = TOP_MARGIN
        
        # === ОБРАБОТКА ТАБЛИЦЫ (КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ) ===
        elif kind == 'table':
            try:
                table_data = _parse_table(block, y_offset, style_cache)
                if table_data:
                    # Проверка, помещается ли таблица
                    if y_offset + table_data['total_height'] > PAGE_HEIGHT - BOTTOM_MARGIN - 20:
                        y_offset = TOP_MARGIN
                    
                    elements.append(make_table_element(
                        x=LEFT_MARGIN, y=y_offset, width=CONTENT_WIDTH, 
                        height=table_data['total_height'], data=table_data['rows']
                    ))
                    y_offset += table_data['total_height'] + 12
                    plain_chunks.append("[Table]")
            except Exception as e:
                print(f"Table parsing error: {e}")

    return {
        "elements": elements,
        "text": "\n".join(plain_chunks),
        "metadata": {
            "page_width": PAGE_WIDTH,
            "page_height": PAGE_HEIGHT,
            "margins": {
                "left": LEFT_MARGIN,
                "right": RIGHT_MARGIN,
                "top": TOP_MARGIN,
                "bottom": BOTTOM_MARGIN
            }
        }
    }

def _build_style_cache(doc: DocumentObject) -> Dict[str, ParagraphFormat]:
    """Кэширование всех стилей документа"""
    cache = {}
    try:
        for style in doc.styles:
            if style.type == 1:  # Paragraph style
                fmt = ParagraphFormat()
                
                style_fmt = style.paragraph_format
                if style_fmt:
                    if style_fmt.left_indent: fmt.left_indent = int(style_fmt.left_indent.pt * PT_TO_PX)
                    if style_fmt.right_indent: fmt.right_indent = int(style_fmt.right_indent.pt * PT_TO_PX)
                    if style_fmt.first_line_indent: fmt.first_line_indent = int(style_fmt.first_line_indent.pt * PT_TO_PX)
                    if style_fmt.space_before: fmt.space_before = int(style_fmt.space_before.pt * PT_TO_PX)
                    if style_fmt.space_after: fmt.space_after = int(style_fmt.space_after.pt * PT_TO_PX)
                    if style_fmt.line_spacing: 
                        fmt.line_spacing = int(style_fmt.line_spacing * PT_TO_PX) or int(fmt.size * 1.15)
                else:
                    fmt.line_spacing = int(fmt.size * 1.15)
                
                if style.font:
                    if style.font.bold: fmt.bold = True
                    if style.font.italic: fmt.italic = True
                    if style.font.size: fmt.size = int(style.font.size.pt)
                    if style.font.name: fmt.font = style.font.name
                
                cache[style.name] = fmt
    except Exception:
        pass
    
    if "Normal" not in cache:
        cache["Normal"] = ParagraphFormat()
    
    return cache

def _process_image(img_bytes, x, y):
    """Обработка изображения"""
    if not img_bytes:
        return None, 0
    
    try:
        pil_img = PILImage.open(io.BytesIO(img_bytes))
        w, h = pil_img.size
        
        max_width = CONTENT_WIDTH - (x - LEFT_MARGIN)
        scale = min(max_width / w, 1.0) if w > max_width else 1.0
        final_w = int(w * scale)
        final_h = int(h * scale)
        
        ext = (pil_img.format or 'png').lower()
        
        if is_likely_signature(final_w, final_h):
            return make_signature_element(x, y, final_w, final_h, img_bytes, ext), final_h
        else:
            return make_image_element(x, y, final_w, final_h, img_bytes, ext), final_h
    except Exception:
        return None, 0

def _iter_all_blocks(doc: DocumentObject):
    """Обход всех блоков документа"""
    section_idx = 0
    for child in doc.element.body.iterchildren():
        if child.tag.endswith('}p'):
            yield ('paragraph', Paragraph(child, doc), section_idx)
        elif child.tag.endswith('}tbl'):
            yield ('table', Table(child, doc), section_idx)
        elif child.tag.endswith('}AlternateContent'):
            for choice in child.iterchildren():
                for inner in choice.iterchildren():
                    if inner.tag.endswith('}p'):
                        yield ('paragraph', Paragraph(inner, doc), section_idx)
                    elif inner.tag.endswith('}tbl'):
                        yield ('table', Table(inner, doc), section_idx)

class ParagraphFormat:
    def __init__(self):
        self.left_indent = 0
        self.right_indent = 0
        self.first_line_indent = 0
        self.space_before = 0
        self.space_after = 0
        self.line_spacing = 16
        self.align = 'left'
        self.bold = False
        self.italic = False
        self.size = 14
        self.font = "Inter"
        self.color = "#000000"
        self.has_border_bottom = False
        self.border_width = 2

def _get_paragraph_format(par: Paragraph, style_cache: Dict[str, ParagraphFormat]) -> ParagraphFormat:
    """Получение форматирования с учетом стилей"""
    fmt = ParagraphFormat()
    
    # === БАЗОВЫЙ СТИЛЬ ===
    if par.style and par.style.name:
        base_fmt = style_cache.get(par.style.name)
        if base_fmt:
            for attr in ['left_indent', 'right_indent', 'first_line_indent', 'space_before', 
                         'space_after', 'line_spacing', 'align', 'bold', 'italic', 'size', 'font']:
                setattr(fmt, attr, getattr(base_fmt, attr))
    
    # === ФОРМАТИРОВАНИЕ ПАРАГРАФА ===
    try:
        par_fmt = par.paragraph_format
        if par_fmt.left_indent is not None: fmt.left_indent = int(par_fmt.left_indent.pt * PT_TO_PX)
        if par_fmt.right_indent is not None: fmt.right_indent = int(par_fmt.right_indent.pt * PT_TO_PX)
        if par_fmt.first_line_indent is not None: fmt.first_line_indent = int(par_fmt.first_line_indent.pt * PT_TO_PX)
        if par_fmt.space_before is not None: fmt.space_before = int(par_fmt.space_before.pt * PT_TO_PX)
        if par_fmt.space_after is not None: fmt.space_after = int(par_fmt.space_after.pt * PT_TO_PX)
        if par_fmt.line_spacing is not None: fmt.line_spacing = int(par_fmt.line_spacing * PT_TO_PX)
        
        if par.alignment is not None:
            align_map = {0: 'left', 1: 'center', 2: 'right', 3: 'justify'}
            fmt.align = align_map.get(par.alignment, 'left')
    except Exception:
        pass
    
    # === RUNS ===
    _apply_runs_format(par, fmt)
    
    # === ГРАНИЦЫ ===
    _check_border_bottom(par, fmt)
    
    return fmt

def _apply_runs_format(par: Paragraph, fmt: ParagraphFormat):
    """Применение форматирования из runs"""
    if not par.runs:
        return
    
    main_run = next((r for r in par.runs if r.text.strip()), par.runs[0])
    
    if main_run.bold: fmt.bold = True
    if main_run.italic: fmt.italic = True
    
    if main_run.font:
        if main_run.font.size: fmt.size = int(main_run.font.size.pt)
        if main_run.font.name: fmt.font = main_run.font.name
        if main_run.font.color and main_run.font.color.rgb:
            fmt.color = f"#{main_run.font.color.rgb}"

def _check_border_bottom(par: Paragraph, fmt: ParagraphFormat):
    """Проверка наличия нижней границы"""
    try:
        pPr = par._element.pPr
        if pPr is None:
            return
        
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is None:
            return
        
        bottom = pBdr.find(qn('w:bottom'))
        if bottom is not None and bottom.get(qn('w:val')) not in (None, 'none', 'nil'):
            fmt.has_border_bottom = True
            sz = bottom.get(qn('w:sz'))
            fmt.border_width = max(1, int((int(sz) or 4) * PT_TO_PX / 12))
    except Exception:
        pass

def _parse_table(table: Table, current_y: int, style_cache: Dict[str, ParagraphFormat]) -> Optional[Dict[str, Any]]:
    """Парсинг таблицы"""
    rows_data = []
    row_heights = []
    col_widths = _get_table_column_widths(table)
    
    for row in table.rows:
        r_cells = []
        max_cell_height = 0
        
        for cell_idx, cell in enumerate(row.cells):
            cell_text = "\n".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
            r_cells.append(cell_text)
            
            cell_w = col_widths[cell_idx] if cell_idx < len(col_widths) else CONTENT_WIDTH // len(row.cells)
            cell_h = _calculate_table_cell_height(cell, cell_w, style_cache)
            
            if cell_h > max_cell_height:
                max_cell_height = cell_h
        
        rows_data.append(r_cells)
        row_heights.append(max_cell_height)
    
    if not rows_data:
        return None
    
    max_cols = max(len(r) for r in rows_data)
    for r in rows_data:
        while len(r) < max_cols:
            r.append("")
    
    return {
        'rows': rows_data,
        'row_heights': row_heights,
        'total_height': sum(row_heights)
    }

def _calculate_table_cell_height(cell: _Cell, width_px: int, style_cache: Dict[str, ParagraphFormat]) -> int:
    """ТОЧНЫЙ расчет высоты ячейки"""
    total_height = 8  # padding top + bottom
    
    for par in cell.paragraphs:
        cell_fmt = _get_paragraph_format(par, style_cache)
        text_content = par.text.strip()
        
        if text_content:
            h = _calculate_text_height_exact(text_content, cell_fmt.size, width_px - 10, cell_fmt.line_spacing)
            total_height += h + cell_fmt.space_before + cell_fmt.space_after
        else:
            total_height += cell_fmt.line_spacing + 4
    
    total_height += 8
    
    return total_height

def _calculate_text_height_exact(text: str, font_size: int, width_px: int, line_spacing: float) -> int:
    """ТОЧНЕЙШАЯ высота текста с переносами"""
    if not text:
        return 0
    
    # line_spacing в px
    line_h = max(int(font_size * 1.15), line_spacing)
    char_w = font_size * 0.55
    
    # Символов в строке
    chars_per_line = max(1, int(width_px // char_w))
    
    total_lines = 0
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            total_lines += 1
            continue
        
        words = line.split()
        current_line_len = 0
        
        for word in words:
            word_len = len(word)
            if word_len > chars_per_line:
                # Длинное слово переносится
                wrapped_lines = math.ceil(word_len / chars_per_line)
                total_lines += wrapped_lines
                current_line_len = word_len % chars_per_line
            elif current_line_len + word_len + 1 <= chars_per_line:
                current_line_len += word_len + 1
            else:
                total_lines += 1
                current_line_len = word_len + 1
        
        total_lines += 1
    
    return total_lines * line_h

def _get_table_column_widths(table: Table) -> List[int]:
    """ТОЧНЫЕ ширины колонок из XML"""
    try:
        tbl_grid = table._element.find(qn('w:tblGrid'))
        if tbl_grid is not None:
            widths = []
            for grid_col in tbl_grid.findall(qn('w:gridCol')):
                w = grid_col.get(qn('w:w'))
                if w:
                    # twip to px: 1 twip = 1/20 pt = 0.0666667 px
                    width_px = int(int(w) * 0.0666667)
                    widths.append(width_px)
            if widths:
                return widths
        
        # Резервный вариант
        if table.rows:
            cell_count = len(table.rows[0].cells)
            if cell_count > 0:
                return [CONTENT_WIDTH // cell_count] * cell_count
        
    except Exception:
        pass
    
    return [CONTENT_WIDTH]

def _has_page_break(par: Paragraph) -> bool:
    """Проверка разрыва страницы"""
    try:
        pPr = par._element.pPr
        if pPr is not None and pPr.find(qn('w:pageBreakBefore')) is not None:
            return True
        
        for run in par.runs:
            brs = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
            for br in brs:
                if br.get(qn('w:type')) == 'page':
                    return True
    except Exception:
        pass
    return False

def _load_document_margins(doc: DocumentObject):
    """Загрузка полей документа"""
    global LEFT_MARGIN, RIGHT_MARGIN, TOP_MARGIN, BOTTOM_MARGIN, CONTENT_WIDTH
    
    try:
        if doc.sections:
            section = doc.sections[0]
            if section.left_margin: LEFT_MARGIN = int(section.left_margin.pt * PT_TO_PX)
            if section.right_margin: RIGHT_MARGIN = int(section.right_margin.pt * PT_TO_PX)
            if section.top_margin: TOP_MARGIN = int(section.top_margin.pt * PT_TO_PX)
            if section.bottom_margin: BOTTOM_MARGIN = int(section.bottom_margin.pt * PT_TO_PX)
            CONTENT_WIDTH = PAGE_WIDTH - (LEFT_MARGIN + RIGHT_MARGIN)
    except Exception:
        pass